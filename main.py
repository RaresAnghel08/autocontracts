from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO
from docx.shared import Inches
import base64
from PIL import Image
from docx2pdf import convert
import shutil
import subprocess
import io, datetime, os, re, copy
import logging
try:
    import pythoncom
except Exception:
    pythoncom = None

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from dotenv import load_dotenv
from supabase import create_client, Client

load_dotenv()

app = Flask(__name__)

# Supabase client
SUPABASE_URL = os.environ.get('SUPABASE_URL', 'https://vddvzvqkdhzyxjtgvgyz.supabase.co')
SUPABASE_KEY = os.environ.get('SUPABASE_KEY', 'sb_publishable_BujADxbZNNVZke-d8vwEuw_lj7kEBsh')
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

@app.route('/test_db')
def test_db():
    try:
        # Test Supabase connection
        response = supabase.table('contract_marshmallow').select('id').limit(1).execute()
        return "Database connection successful"
    except Exception as e:
        return f"Database connection failed: {e}"


@app.route('/download/<child_folder>/<filename>')
def download_file(child_folder, filename):
    # basic validation to prevent directory traversal
    if '/' in child_folder or '..' in child_folder or '/' in filename or '..' in filename:
        return "Invalid path", 400
    safe_base = os.path.join(os.getcwd(), 'temp')
    file_path = os.path.join(safe_base, child_folder, filename)
    if not os.path.exists(file_path):
        return "File not found", 404
    return send_file(file_path, as_attachment=True)

# Note: HTML moved to templates/form.html and static/styles.css

@app.route('/')
def form():
    return render_template('form.html')

@app.route('/generate', methods=['POST'])
def generate_docx():
    data = request.form.to_dict()

    # Generare automată număr și dată contract
    today = datetime.datetime.now()
    kindergarten = data.get('kindergarten', 'ariel')
    prefix = 'MARSHMALLOW' if kindergarten == 'marshmallow' else 'ARIEL'
    data['data_contract'] = today.strftime('%d.%m.%Y')
    data['numar_contract'] = f"{prefix}-{today.strftime('%Y%m%d-%H%M%S')}"

    # Save to database using Supabase
    kindergarten = data.get('kindergarten', 'ariel')
    contract_data = {
        'nume_copil': data.get('nume_copil'),
        'prenume_copil': data.get('prenume_copil'),
        'cnp_copil': data.get('cnp_copil'),
        'data_nasterii_copil': data.get('data_nasterii_copil'),
        'adresa_copil': data.get('adresa_copil'),
        'nume_mama': data.get('nume_mama'),
        'cnp_mama': data.get('cnp_mama'),
        'serie_buletin_mama': data.get('serie_buletin_mama'),
        'numar_buletin_mama': data.get('numar_buletin_mama'),
        'eliberat_de_mama': data.get('eliberat_de_mama'),
        'data_eliberarii_mama': data.get('data_eliberarii_mama'),
        'adresa_mama': data.get('adresa_mama'),
        'email_mama': data.get('email_mama'),
        'telefon_mama': data.get('telefon_mama'),
        'nume_tata': data.get('nume_tata'),
        'cnp_tata': data.get('cnp_tata'),
        'serie_buletin_tata': data.get('serie_buletin_tata'),
        'numar_buletin_tata': data.get('numar_buletin_tata'),
        'eliberat_de_tata': data.get('eliberat_de_tata'),
        'data_eliberarii_tata': data.get('data_eliberarii_tata'),
        'adresa_tata': data.get('adresa_tata'),
        'email_tata': data.get('email_tata'),
        'telefon_tata': data.get('telefon_tata'),
        'telefon_bunici': data.get('telefon_bunici'),
        'tel_urgenta': data.get('tel_urgenta'),
        'probleme_medicale': data.get('probleme_medicale'),
        'alergii': data.get('alergii'),
        'persoane_autorizate': data.get('persoane_autorizate'),
        'alte_observatii': data.get('alte_observatii'),
        'gdpr_1': data.get('1') == 'da',
        'gdpr_2': data.get('2') == 'da',
        'supraveghere_video': data.get('3'),
        'program': data.get('program'),
        'grupa': data.get('5'),
        'signature_data': data.get('signature_data'),
        'data_contract': data['data_contract'],
        'numar_contract': data['numar_contract']
    }
    
    try:
        if kindergarten == 'marshmallow':
            supabase.table('contract_marshmallow').insert(contract_data).execute()
        else:
            supabase.table('contract_ariel').insert(contract_data).execute()
    except Exception as e:
        print(f"Database error: {e}")
        # Continue with PDF generation even if database save fails

    # Creăm un folder temporar sigur în proiect per copil
    base_temp = os.path.join(os.getcwd(), 'temp')
    os.makedirs(base_temp, exist_ok=True)

    # Construim numele folderului din nume + prenume copil și îl sanitizăm
    child_raw = f"{data.get('nume_copil','').strip()}_{data.get('prenume_copil','').strip()}".strip('_ ')
    def _sanitize(s):
        s = s.strip()
        s = re.sub(r'[^A-Za-z0-9_-]', '_', s)
        s = re.sub(r'_+', '_', s)
        return s or 'copil'

    child_name = _sanitize(child_raw)

    # Dacă folderul există deja, adăugăm sufix _1, _2, ...
    candidate = os.path.join(base_temp, child_name)
    count = 1
    while os.path.exists(candidate):
        candidate = os.path.join(base_temp, f"{child_name}_{count}")
        count += 1

    temp_dir = candidate
    os.makedirs(temp_dir, exist_ok=True)

    # Move any stray files from the base temp folder into this child's folder
    try:
        for entry in os.listdir(base_temp):
            entry_path = os.path.join(base_temp, entry)
            # skip directories (including other child folders)
            if os.path.isfile(entry_path):
                # move files (e.g., leftover .docx/.pdf/.png) into the child folder
                shutil.move(entry_path, os.path.join(temp_dir, entry))
    except Exception:
        # don't fail the whole request if cleanup/move can't run
        pass
    signature_path = os.path.join(temp_dir, 'signature.png')

    # Salvăm semnătura ca imagine
    signature_data = data['signature_data'].split(',')[1]
    signature_bytes = base64.b64decode(signature_data)
    signature_img = Image.open(io.BytesIO(signature_bytes))
    signature_img.save(signature_path)

    # Template-uri DOCX (use project-relative paths)
    base_dir = os.path.dirname(os.path.abspath(__file__))
    templates_dir = os.path.join(base_dir, 'templates')
    contracts = [
        (os.path.join(templates_dir, 'educational.docx'), "contract_educational_completat.docx"),
        (os.path.join(templates_dir, 'catering.docx'), "contract_catering_completat.docx")
    ]

    generated_pdfs = []
    # compile a regex to catch {{semnatura}} with optional spaces and case-insensitive
    sig_pattern = re.compile(r"\{\{\s*semnatura\s*\}\}", flags=re.IGNORECASE)
    for template_path, output_name in contracts:
        doc = Document(template_path)
        signature_inserted = False

        def replace_placeholders_in_paragraph(p):
            nonlocal signature_inserted
            # replace normal placeholders from form data
            # skip numeric placeholders handled specially (e.g. '3', '4', '5')
            for key, val in data.items():
                if key not in ('signature_data', '3', '4', '5'):
                    if f'{{{{{key}}}}}' in p.text:
                        p.text = p.text.replace(f'{{{{{key}}}}}', val)

            # special for {{4}} program selection: use checkbox characters so boxes remain visible
            if '{{4}}' in p.text:
                prog = data.get('program', '')
                # use checked/unchecked box unicode characters
                checked = '☑'
                unchecked = '☐'
                if prog == 'normal':
                    # first {{4}} -> checked, second -> unchecked
                    p.text = p.text.replace('{{4}}', checked, 1).replace('{{4}}', unchecked, 1)
                elif prog == 'prelungit':
                    # first -> unchecked, second -> checked
                    p.text = p.text.replace('{{4}}', unchecked, 1).replace('{{4}}', checked, 1)
                else:
                    p.text = p.text.replace('{{4}}', unchecked)

            # special for {{5}} group selection: render checked/unchecked for four group options
            if '{{5}}' in p.text:
                grp = data.get('5', '').lower()
                checked = '☑'
                unchecked = '☐'
                # expected four placeholders in template for the four groups
                if p.text.count('{{5}}') >= 4:
                    order = ['mica', 'mica_b', 'mijlocie', 'mare']
                    new_text = p.text
                    for opt in order:
                        if grp == opt:
                            new_text = new_text.replace('{{5}}', checked, 1)
                        else:
                            new_text = new_text.replace('{{5}}', unchecked, 1)
                    p.text = new_text
                else:
                    # fallback: single placeholder -> checked if any group chosen
                    p.text = p.text.replace('{{5}}', checked if grp else unchecked)

            # special for {{3}} consent: use checkbox characters and preserve labels
            if '{{3}}' in p.text:
                consent = data.get('3', '').lower()
                checked = '☑'
                unchecked = '☐'
                is_agree = consent in ('da', 'sunt', 'sunt de acord', 'true', 'on', 'yes')
                # If template contains two placeholders (one per option), mark first/second appropriately
                if p.text.count('{{3}}') >= 2:
                    if is_agree:
                        # first -> checked, second -> unchecked
                        p.text = p.text.replace('{{3}}', checked, 1).replace('{{3}}', unchecked, 1)
                    else:
                        # first -> unchecked, second -> checked
                        p.text = p.text.replace('{{3}}', unchecked, 1).replace('{{3}}', checked, 1)
                else:
                    # single placeholder: replace with checked/unchecked
                    p.text = p.text.replace('{{3}}', checked if is_agree else unchecked)

            # check for signature placeholder variants
            if sig_pattern.search(p.text):
                # remove placeholder
                p.text = sig_pattern.sub('', p.text)
                try:
                    # add picture which creates a new paragraph with the image
                    width = Inches(2)
                    doc.add_picture(signature_path, width=width)
                    pic_para = doc.paragraphs[-1]
                    # get the drawing element from the picture run
                    pic_run = pic_para.runs[0] if pic_para.runs else None
                    drawing_elem = None
                    if pic_run is not None:
                        for child in pic_run._r:
                            # look for drawing element (namespace-aware)
                            if 'drawing' in child.tag:
                                drawing_elem = child
                                break
                    if drawing_elem is not None:
                        # create a new run in the target paragraph and insert the drawing
                        new_run = p.add_run()
                        new_run._r.append(copy.deepcopy(drawing_elem))
                        # move the new run to the start of the paragraph so image appears before text
                        p._p.insert(0, new_run._r)
                        # remove the temporary picture paragraph
                        try:
                            pic_para._p.getparent().remove(pic_para._p)
                        except Exception:
                            pass
                        signature_inserted = True
                except Exception:
                    pass

        # process top-level paragraphs
        for p in doc.paragraphs:
            replace_placeholders_in_paragraph(p)

        # process paragraphs inside table cells as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholders_in_paragraph(p)

        # process textboxes in shapes
        for shape in doc.inline_shapes:
            if hasattr(shape, 'has_text_frame') and shape.has_text_frame:
                for p in shape.text_frame.paragraphs:
                    replace_placeholders_in_paragraph(p)

        # If no {{semnatura}} was present/inserted, add signature block at end
        if not signature_inserted:
            doc.add_paragraph(f"\nSemnătură părinte:\n{data.get('nume_mama', '')} / {data.get('nume_tata', '')}")
            try:
                width = Inches(2)
                doc.add_picture(signature_path, width=width)
            except Exception:
                pass
        output_path = os.path.join(temp_dir, output_name)
        doc.save(output_path)

        # Convert DOCX to PDF (try docx2pdf, then LibreOffice/soffice fallback)
        pdf_name = os.path.splitext(output_name)[0] + '.pdf'
        pdf_path = os.path.join(temp_dir, pdf_name)

        def try_convert_to_pdf(input_docx, output_pdf):
            # Try docx2pdf (MS Word COM on Windows)
            try:
                if pythoncom is not None:
                    try:
                        pythoncom.CoInitialize()
                    except Exception:
                        pass
                    convert(input_docx, output_pdf)
                    return True
                else:
                    # pythoncom not available (non-Windows); skip docx2pdf
                    logging.info("pythoncom not available; skipping docx2pdf (Word COM)")
            except Exception as e:
                logging.error(f"docx2pdf failed for {input_docx}: {e}")
                pass
            # Try LibreOffice/soffice headless conversion
            soffice = shutil.which('soffice') or shutil.which('libreoffice')
            if soffice:
                try:
                    outdir = os.path.dirname(output_pdf)
                    subprocess.run([soffice, '--headless', '--convert-to', 'pdf', '--outdir', outdir, input_docx], check=True)
                    return os.path.exists(output_pdf)
                except Exception as e:
                    logging.error(f"soffice failed for {input_docx}: {e}")
                    return False
            return False

        converted = try_convert_to_pdf(output_path, pdf_path)
        if converted:
            generated_pdfs.append(pdf_name)
        else:
            # conversion failed; include DOCX as fallback
            generated_pdfs.append(os.path.basename(output_path))

    # Prepare data for the download template: pairs of (display name, link)
    final_child_folder = os.path.basename(temp_dir)
    files_for_template = [(name, f"/download/{final_child_folder}/{name}") for name in generated_pdfs]
    # failed conversions are those that are not PDF files
    failed = [name for name in generated_pdfs if not name.lower().endswith('.pdf')]

    # Render the nicer download page (templates/download.html)
    return render_template('download.html', files=files_for_template, failed=failed)

@app.route('/admin')
def admin():
    search = request.args.get('search', '')
    kindergarten = request.args.get('kindergarten', 'ariel')
    
    try:
        if kindergarten == 'marshmallow':
            table_name = 'contract_marshmallow'
        else:
            table_name = 'contract_ariel'
        
        # Build query
        query = supabase.table(table_name).select('*').order('created_at', desc=True)
        
        if search:
            # Supabase doesn't have direct ILIKE support in Python client, so we'll filter in Python
            response = query.execute()
            contracts = response.data
            # Filter by search term
            contracts = [
                contract for contract in contracts
                if search.lower() in (contract.get('nume_copil') or '').lower() or
                   search.lower() in (contract.get('prenume_copil') or '').lower()
            ]
        else:
            response = query.execute()
            contracts = response.data
            
    except Exception as e:
        print(f"Database error: {e}")
        contracts = []
    
    return render_template('admin.html', contracts=contracts, search=search, kindergarten=kindergarten)

@app.route('/generate_pdf/<kindergarten>/<int:contract_id>/<contract_type>')
def generate_pdf(kindergarten, contract_id, contract_type):
    try:
        if kindergarten == 'marshmallow':
            table_name = 'contract_marshmallow'
        else:
            table_name = 'contract_ariel'
        
        response = supabase.table(table_name).select('*').eq('id', contract_id).execute()
        contract_data = response.data[0] if response.data else None
        
        if not contract_data:
            return "Contract not found", 404
            
    except Exception as e:
        print(f"Database error: {e}")
        return "Database error", 500
    
    # Convert contract data to the format expected by the PDF generation
    data = {
        'nume_copil': contract_data.get('nume_copil'),
        'prenume_copil': contract_data.get('prenume_copil'),
        'cnp_copil': contract_data.get('cnp_copil'),
        'data_nasterii_copil': contract_data.get('data_nasterii_copil'),
        'adresa_copil': contract_data.get('adresa_copil'),
        'nume_mama': contract_data.get('nume_mama'),
        'cnp_mama': contract_data.get('cnp_mama'),
        'serie_buletin_mama': contract_data.get('serie_buletin_mama'),
        'numar_buletin_mama': contract_data.get('numar_buletin_mama'),
        'eliberat_de_mama': contract_data.get('eliberat_de_mama'),
        'data_eliberarii_mama': contract_data.get('data_eliberarii_mama'),
        'adresa_mama': contract_data.get('adresa_mama'),
        'email_mama': contract_data.get('email_mama'),
        'telefon_mama': contract_data.get('telefon_mama'),
        'nume_tata': contract_data.get('nume_tata'),
        'cnp_tata': contract_data.get('cnp_tata'),
        'serie_buletin_tata': contract_data.get('serie_buletin_tata'),
        'numar_buletin_tata': contract_data.get('numar_buletin_tata'),
        'eliberat_de_tata': contract_data.get('eliberat_de_tata'),
        'data_eliberarii_tata': contract_data.get('data_eliberarii_tata'),
        'adresa_tata': contract_data.get('adresa_tata'),
        'email_tata': contract_data.get('email_tata'),
        'telefon_tata': contract_data.get('telefon_tata'),
        'telefon_bunici': contract_data.get('telefon_bunici'),
        'tel_urgenta': contract_data.get('tel_urgenta'),
        'probleme_medicale': contract_data.get('probleme_medicale'),
        'alergii': contract_data.get('alergii'),
        'persoane_autorizate': contract_data.get('persoane_autorizate'),
        'alte_observatii': contract_data.get('alte_observatii'),
        '1': 'da' if contract_data.get('gdpr_1') else '',
        '2': 'da' if contract_data.get('gdpr_2') else '',
        '3': contract_data.get('supraveghere_video'),
        'program': contract_data.get('program'),
        '5': contract_data.get('grupa'),
        'signature_data': contract_data.get('signature_data'),
        'data_contract': contract_data.get('data_contract'),
        'numar_contract': contract_data.get('numar_contract')
    }
    
    # Now generate the PDF using the same logic as generate_docx
    # But instead of returning the download page, return the PDF directly
    
    # Creăm un folder temporar sigur
    base_temp = os.path.join(os.getcwd(), 'temp')
    os.makedirs(base_temp, exist_ok=True)
    
    temp_dir = os.path.join(base_temp, f"pdf_{contract_id}")
    os.makedirs(temp_dir, exist_ok=True)
    
    signature_path = os.path.join(temp_dir, 'signature.png')
    
    # Salvăm semnătura ca imagine
    if contract_data.get('signature_data'):
        signature_data = contract_data.get('signature_data').split(',')[1]
        signature_bytes = base64.b64decode(signature_data)
        signature_img = Image.open(io.BytesIO(signature_bytes))
        signature_img.save(signature_path)
    
    # Template-uri DOCX - select based on contract type
    base_dir = os.path.dirname(os.path.abspath(__file__))
    templates_dir = os.path.join(base_dir, 'templates')
    
    if contract_type == 'educational':
        contracts_templates = [
            (os.path.join(templates_dir, 'educational.docx'), f"contract_educational_{contract_data.get('nume_copil', 'copil')}_{contract_data.get('prenume_copil', '')}.pdf")
        ]
    elif contract_type == 'catering':
        contracts_templates = [
            (os.path.join(templates_dir, 'catering.docx'), f"contract_catering_{contract_data.get('nume_copil', 'copil')}_{contract_data.get('prenume_copil', '')}.pdf")
        ]
    else:
        # Default to both if no specific type
        contracts_templates = [
            (os.path.join(templates_dir, 'educational.docx'), f"contract_educational_{contract_data.get('nume_copil', 'copil')}_{contract_data.get('prenume_copil', '')}.pdf"),
            (os.path.join(templates_dir, 'catering.docx'), f"contract_catering_{contract_data.get('nume_copil', 'copil')}_{contract_data.get('prenume_copil', '')}.pdf")
        ]
    
    generated_pdfs = []
    sig_pattern = re.compile(r"\{\{\s*semnatura\s*\}\}", flags=re.IGNORECASE)
    
    for template_path, output_name in contracts_templates:
        doc = Document(template_path)
        signature_inserted = False
        
        def replace_placeholders_in_paragraph(p):
            nonlocal signature_inserted
            for key, val in data.items():
                if key not in ('signature_data', '3', '4', '5'):
                    if f'{{{{{key}}}}}' in p.text:
                        p.text = p.text.replace(f'{{{{{key}}}}}', str(val) if val else '')
            
            # Special handling for checkboxes and signature as before
            # ... (same logic as in generate_docx)
            
            if '{{4}}' in p.text:
                prog = data.get('program', '')
                checked = '☑'
                unchecked = '☐'
                if prog == 'normal':
                    p.text = p.text.replace('{{4}}', checked, 1).replace('{{4}}', unchecked, 1)
                elif prog == 'prelungit':
                    p.text = p.text.replace('{{4}}', unchecked, 1).replace('{{4}}', checked, 1)
                else:
                    p.text = p.text.replace('{{4}}', unchecked)
            
            if '{{5}}' in p.text:
                grp = data.get('5', '').lower()
                checked = '☑'
                unchecked = '☐'
                if p.text.count('{{5}}') >= 4:
                    order = ['mica', 'mica_b', 'mijlocie', 'mare']
                    new_text = p.text
                    for opt in order:
                        if grp == opt:
                            new_text = new_text.replace('{{5}}', checked, 1)
                        else:
                            new_text = new_text.replace('{{5}}', unchecked, 1)
                    p.text = new_text
                else:
                    p.text = p.text.replace('{{5}}', checked if grp else unchecked)
            
            if '{{3}}' in p.text:
                consent = data.get('3', '').lower()
                checked = '☑'
                unchecked = '☐'
                is_agree = consent in ('da', 'sunt', 'sunt de acord', 'true', 'on', 'yes')
                if p.text.count('{{3}}') >= 2:
                    if is_agree:
                        p.text = p.text.replace('{{3}}', checked, 1).replace('{{3}}', unchecked, 1)
                    else:
                        p.text = p.text.replace('{{3}}', unchecked, 1).replace('{{3}}', checked, 1)
                else:
                    p.text = p.text.replace('{{3}}', checked if is_agree else unchecked)
            
            if sig_pattern.search(p.text):
                p.text = sig_pattern.sub('', p.text)
                try:
                    width = Inches(2)
                    doc.add_picture(signature_path, width=width)
                    pic_para = doc.paragraphs[-1]
                    pic_run = pic_para.runs[0] if pic_para.runs else None
                    drawing_elem = None
                    if pic_run is not None:
                        for child in pic_run._r:
                            if 'drawing' in child.tag:
                                drawing_elem = child
                                break
                    if drawing_elem is not None:
                        new_run = p.add_run()
                        new_run._r.append(copy.deepcopy(drawing_elem))
                        p._p.insert(0, new_run._r)
                        try:
                            pic_para._p.getparent().remove(pic_para._p)
                        except Exception:
                            pass
                        signature_inserted = True
                except Exception:
                    pass
        
        for p in doc.paragraphs:
            replace_placeholders_in_paragraph(p)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholders_in_paragraph(p)
        
        for shape in doc.inline_shapes:
            if hasattr(shape, 'has_text_frame') and shape.has_text_frame:
                for p in shape.text_frame.paragraphs:
                    replace_placeholders_in_paragraph(p)
        
        if not signature_inserted and contract_data.get('signature_data'):
            doc.add_paragraph(f"\nSemnătură părinte:\n{contract_data.get('nume_mama')} / {contract_data.get('nume_tata')}")
            try:
                width = Inches(2)
                doc.add_picture(signature_path, width=width)
            except Exception:
                pass
        
        output_path = os.path.join(temp_dir, output_name.replace('.pdf', '.docx'))
        doc.save(output_path)
        
        pdf_path = os.path.join(temp_dir, output_name)
        
        def try_convert_to_pdf(input_docx, output_pdf):
            # Try docx2pdf
            try:
                if pythoncom is not None:
                    try:
                        pythoncom.CoInitialize()
                    except Exception:
                        pass
                    convert(input_docx, output_pdf)
                    return True
            except Exception:
                pass
            
            # Try LibreOffice
            soffice = shutil.which('soffice') or shutil.which('libreoffice')
            if soffice:
                try:
                    outdir = os.path.dirname(output_pdf)
                    subprocess.run([soffice, '--headless', '--convert-to', 'pdf', '--outdir', outdir, input_docx], check=True)
                    return os.path.exists(output_pdf)
                except Exception:
                    return False
            return False
        
        if try_convert_to_pdf(output_path, pdf_path):
            generated_pdfs.append(pdf_path)
    
    if generated_pdfs:
        # Return the generated PDF with proper filename
        pdf_filename = contracts_templates[0][1]  # Get the filename from the template
        return send_file(generated_pdfs[0], as_attachment=True, download_name=pdf_filename)
    else:
        return "PDF generation failed", 500

if __name__ == '__main__':
    app.run(debug=True)